"""
# My first app
Here's our first attempt at using data to create a table:
"""

import streamlit as st
import pandas as pd
import PyPDF2 
import re
import docx

doc = docx.Document()

#PDF Pages
pdf_read = PyPDF2.PdfReader("C:/Users/Nikhil Chandra/Documents/Repos/StreamLit_Practice/154_Alcock_Road_Elimbah_QLD_4516_07-01-2025.pdf")
page_1 = pdf_read.pages[0].extract_text()
page_2 = pdf_read.pages[1].extract_text()
page_3 = pdf_read.pages[2].extract_text()
adjusted_text_1 = page_1.replace(',', '')
adjusted_text_2 = page_2.replace(',', '')
adjusted_text_3 = page_3.replace(',', '')

# Get Address
adjusted_text_split_copyright = adjusted_text_1.split("Copyright")[0]
address_split = re.search("^.*?(=-|[0-9][0-9][0-9][0-9])",adjusted_text_split_copyright)
address_found = address_split.group(0)
doc.add_heading(address_found)

#Get Prices
doc_x_table = doc.add_table(rows = 3, cols = 2)
first_row = doc_x_table.rows[0].cells
first_row[0].text = ''

found_prices_p1 = re.search("(Sold for\$[0-9]+)|(Sold for \$[0-9]+)",adjusted_text_1)
if (found_prices_p1 is not None):
    print(found_prices_p1)

found_prices_p2 = re.search("(Sold for\$[0-9]+)|(Sold for \$[0-9]+)",adjusted_text_2)
if (found_prices_p2 is not None):
    print(len(found_prices_p2))


def extract_prices(text_to_extract = adjusted_text_3):
    found_prices = re.findall("(Sold for\$[0-9]+)|"
                             "(Sold for \$[0-9]+)|"
                             "(Transferred for \$[0-9]+)"
                            #  "([0-9]+ months and [0-9]+ days ago)|"
                            #  "([0-9]+ years and [0-9]+ months ago)|"
                            #  "([0-9]+ years, [0-9]+ months ago)"
                             ,text_to_extract)
    found_prices_extracted = []
    if (found_prices is not None):
        for i in range(0,len(found_prices)):
            test = [a for a in found_prices[i] if a != '']
            found_prices_extracted.append(test)
            print(test)
            print(i)
    return found_prices_extracted

prices_p1 = extract_prices(text_to_extract = adjusted_text_3)
print(prices_p1)

doc.save(address_found + ".docx")
